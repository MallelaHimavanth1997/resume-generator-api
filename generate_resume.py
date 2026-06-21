#!/usr/bin/env python3
"""
generate_resume.py — Builds a .docx resume matching Himavanth Mallela's
master template (bordered page, Times New Roman, checkmark-bullet summary,
skills table, per-client experience blocks, cert badges) from a structured
JSON input.

Designed to be called by n8n (or any script) as the final step after Claude
tailors resume content: Claude outputs structured JSON -> this script ->
formatted .docx matching the original template exactly.

Usage:
    python3 generate_resume.py input.json output.docx

See sample_input.json for the expected JSON shape.
"""

import sys
import json
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
ASSETS_DIR = os.path.join(SCRIPT_DIR, "assets")
FONT = "Times New Roman"
BODY_SIZE = Pt(11)
NAME_SIZE = Pt(16)


# ---------- low-level XML helpers ----------

def set_page_border(section):
    sectPr = section._sectPr
    pgBorders = OxmlElement("w:pgBorders")
    pgBorders.set(qn("w:offsetFrom"), "page")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "24")
        el.set(qn("w:color"), "auto")
        pgBorders.append(el)
    # CT_SectPr schema order: ...pgSz, pgMar, paperSrc, pgBorders, lnNumType, pgNumType, cols, ...
    pgMar = sectPr.find(qn("w:pgMar"))
    if pgMar is not None:
        pgMar.addnext(pgBorders)
    else:
        sectPr.insert(0, pgBorders)


def set_cell_borders_none(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    # CT_TcPr schema order: tcW, gridSpan, hMerge/vMerge, tcBorders, shd, ...
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is not None:
        tcW.addnext(borders)
    else:
        tcPr.insert(0, borders)


def set_table_borders(table, color="000000", sz="4"):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    # CT_TblPrBase schema order: tblStyle, tblpPr, tblOverlap, bidiVisual, tblStyleRowBandSize,
    # tblStyleColBandSize, tblW, jc, tblCellSpacing, tblInd, tblBorders, shd, tblLayout, tblCellMar, ...
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is not None:
        tblW.addnext(borders)
    else:
        tblPr.insert(0, borders)


def add_bottom_rule(paragraph, size=18, color="000000"):
    """Adds a horizontal rule via paragraph bottom border (used under header)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    # pBdr must be inserted in correct schema position within pPr (after pStyle/keepNext/etc,
    # before spacing/ind/jc). Safest: insert at index 0 unless pStyle present.
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is not None:
        pStyle.addnext(pBdr)
    else:
        pPr.insert(0, pBdr)


def set_run_font(run, bold=False, size=BODY_SIZE, color=None):
    run.font.name = FONT
    run.font.size = size
    run.font.bold = bold
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)
    rFonts.set(qn("w:cs"), FONT)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_run(paragraph, text, bold=False, size=BODY_SIZE, underline=False, color=None, hyperlink_color=None):
    run = paragraph.add_run(text)
    set_run_font(run, bold=bold, size=size, color=color)
    run.font.underline = underline
    return run


def add_rich_paragraph(doc_or_cell, segments, bullet_char=None, indent_left=Inches(0.25),
                        hanging=Inches(0.25), space_after=Pt(6), space_before=Pt(0),
                        alignment=None):
    """segments: list of (text, bold) tuples. Renders bold spans inline like '**word**' parsing result."""
    p = doc_or_cell.add_paragraph()
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    if alignment:
        p.alignment = alignment
    if bullet_char:
        p.paragraph_format.left_indent = indent_left
        p.paragraph_format.first_line_indent = -hanging
        add_run(p, bullet_char + "  ", bold=False)
    for text, bold in segments:
        add_run(p, text, bold=bold)
    return p


def parse_bold_markup(text):
    """Converts '**bold**' inline markup into (text, bold) segments."""
    segments = []
    parts = text.split("**")
    for i, part in enumerate(parts):
        if part == "":
            continue
        segments.append((part, i % 2 == 1))
    return segments if segments else [(text, False)]


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = add_run(p, text, bold=True)
    run.font.underline = True
    return p


# ---------- section builders ----------

def build_header(doc, data):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    content_width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
    right_w = Inches(2.7)
    left_w = content_width - right_w
    table.columns[0].width = left_w
    table.columns[1].width = right_w

    left_cell, right_cell = table.rows[0].cells
    left_cell.width = left_w
    right_cell.width = right_w
    set_cell_borders_none(left_cell)
    set_cell_borders_none(right_cell)

    left_cell.paragraphs[0].text = ""
    p = left_cell.paragraphs[0]
    add_run(p, data["name"], bold=True, size=NAME_SIZE)

    title_p = left_cell.add_paragraph()
    add_run(title_p, data.get("title", ""), bold=False)

    phone_p = left_cell.add_paragraph()
    add_run(phone_p, "Phone: ", bold=True)
    add_run(phone_p, data.get("phone", ""), bold=False)

    email_p = left_cell.add_paragraph()
    add_run(email_p, "Email: ", bold=True)
    email_run = add_run(email_p, data.get("email", ""), bold=False, color="0563C1")
    email_run.font.underline = True

    # Cert badges in right cell, side by side via a nested borderless table
    badges = data.get("cert_badge_images", [])
    if badges:
        right_cell.paragraphs[0].text = ""
        inner = right_cell.add_table(rows=1, cols=len(badges))
        inner.autofit = False
        badge_col_w = Inches(2.7 / max(len(badges), 1))
        set_cell_borders_none_for_table(inner)
        for i, badge_file in enumerate(badges):
            cell = inner.rows[0].cells[i]
            cell.width = badge_col_w
            set_cell_borders_none(cell)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell.paragraphs[0].add_run()
            img_path = os.path.join(ASSETS_DIR, badge_file)
            if os.path.exists(img_path):
                run.add_picture(img_path, width=Inches(0.78))

    # bottom rule under header
    rule_p = doc.add_paragraph()
    rule_p.paragraph_format.space_before = Pt(4)
    rule_p.paragraph_format.space_after = Pt(8)
    add_bottom_rule(rule_p, size=18)


def set_cell_borders_none_for_table(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is not None:
        tblW.addnext(borders)
    else:
        tblPr.insert(0, borders)


def build_summary(doc, data):
    add_heading(doc, "Professional Summary:")
    for line in data.get("summary_bullets", []):
        segments = parse_bold_markup(line)
        add_rich_paragraph(doc, segments, bullet_char="\u27a4", space_after=Pt(8))


def build_skills_table(doc, data):
    add_heading(doc, "Technical Skills:")
    skills = data.get("technical_skills", [])
    if not skills:
        return
    table = doc.add_table(rows=0, cols=2)
    set_table_borders(table)
    content_width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
    col1_w = Inches(1.7)
    col2_w = content_width - col1_w
    for category, values in skills:
        row = table.add_row()
        c1, c2 = row.cells
        c1.width = col1_w
        c2.width = col2_w
        c1.paragraphs[0].text = ""
        add_run(c1.paragraphs[0], category, bold=True)
        c2.paragraphs[0].text = ""
        add_run(c2.paragraphs[0], values, bold=False)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


def build_experience(doc, data):
    add_heading(doc, "Professional Experience:")
    for job in data.get("experience", []):
        client_p = doc.add_paragraph()
        client_p.paragraph_format.space_before = Pt(8)
        client_p.paragraph_format.space_after = Pt(0)
        add_run(client_p, f"Client: {job['client']}", bold=True)
        tab_stops = client_p.paragraph_format.tab_stops
        content_width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
        tab_stops.add_tab_stop(content_width, alignment=2)  # right-aligned tab
        add_run(client_p, "\t" + job.get("dates", ""), bold=True)

        role_p = doc.add_paragraph()
        role_p.paragraph_format.space_after = Pt(8)
        add_run(role_p, f"Role: {job['role']}", bold=True)

        resp_heading = doc.add_paragraph()
        resp_heading.paragraph_format.space_after = Pt(4)
        run = add_run(resp_heading, "Responsibilities:", bold=True)
        run.font.underline = True

        for bullet in job.get("responsibilities", []):
            segments = parse_bold_markup(bullet)
            add_rich_paragraph(doc, segments, bullet_char="\u2022", space_after=Pt(6))

        if job.get("environment"):
            env_p = doc.add_paragraph()
            env_p.paragraph_format.space_before = Pt(6)
            env_p.paragraph_format.space_after = Pt(10)
            add_run(env_p, "Environment: ", bold=True)
            add_run(env_p, job["environment"], bold=False)


def build_education(doc, data):
    edu = data.get("education")
    if not edu:
        return
    add_heading(doc, "EDUCATION:")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    segments = parse_bold_markup(edu)
    for text, bold in segments:
        add_run(p, text, bold=bold)


def build_projects(doc, data):
    projects = data.get("projects", [])
    if not projects:
        return
    add_heading(doc, "PROJECTS:")
    for proj in projects:
        title_p = doc.add_paragraph()
        title_p.paragraph_format.space_after = Pt(2)
        add_run(title_p, proj["title"], bold=True)
        for bullet in proj.get("bullets", []):
            segments = parse_bold_markup(bullet)
            add_rich_paragraph(doc, segments, bullet_char="\u2022",
                                space_after=Pt(2), indent_left=Inches(0.15), hanging=Inches(0.15))
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(2)


def build_certifications(doc, data):
    certs = data.get("certifications", [])
    if not certs:
        return
    add_heading(doc, "CERTIFICATIONS:")
    for cert in certs:
        name_p = doc.add_paragraph()
        name_p.paragraph_format.space_before = Pt(8)
        name_p.paragraph_format.space_after = Pt(2)
        add_run(name_p, cert["name"], bold=True)
        if cert.get("description"):
            desc_p = doc.add_paragraph()
            desc_p.paragraph_format.space_after = Pt(2)
            segments = parse_bold_markup(cert["description"])
            for text, bold in segments:
                add_run(desc_p, text, bold=bold)


# ---------- main ----------

def build_resume(data, output_path):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    set_page_border(section)

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = BODY_SIZE

    # Fix python-docx default template's incomplete zoom element (missing required percent attr)
    zoom = doc.settings.element.find(qn("w:zoom"))
    if zoom is not None:
        zoom.set(qn("w:percent"), "100")

    build_header(doc, data)
    build_summary(doc, data)
    build_skills_table(doc, data)
    build_experience(doc, data)
    build_education(doc, data)
    build_projects(doc, data)
    build_certifications(doc, data)

    doc.save(output_path)
    return output_path


def main():
    if len(sys.argv) != 3:
        print("Usage: python3 generate_resume.py input.json output.docx")
        sys.exit(1)
    input_path, output_path = sys.argv[1], sys.argv[2]
    with open(input_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    build_resume(data, output_path)
    print(f"Saved: {output_path}")


if __name__ == "__main__":
    main()
